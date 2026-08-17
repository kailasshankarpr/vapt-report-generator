from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.exporters.base_exporter import BaseExporter
from src.models.vulnerability import ScanReport, Severity


class DOCXExporter(BaseExporter):
    """Generates styled editable Microsoft Word (.docx) reports"""

    def export(self, report: ScanReport, output_path: str) -> str:
        doc = Document()

        # Document Title
        title_p = doc.add_heading(level=0)
        run = title_p.add_run(report.title)
        run.font.color.rgb = RGBColor(15, 23, 42)  # Slate-900

        # Subtitle Info
        meta_p = doc.add_paragraph()
        meta_p.add_run(f"Client: {report.client_name}   |   Project: {report.project_name}   |   Date: {datetime.now().strftime('%Y-%m-%d')}\n")

        # Executive Dashboard Table
        doc.add_heading("Executive Dashboard Summary", level=1)
        dash_table = doc.add_table(rows=2, cols=4)
        dash_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Security Grade", "Total Vulnerabilities", "CVSS Risk Score", "Compliance Avg"]
        avg_comp = round(sum(d['compliance_score'] for d in report.compliance_summary.values()) / max(1, len(report.compliance_summary)), 1)
        vals = [report.security_health_grade, str(report.total_vulnerabilities), f"{report.overall_risk_score} / 10", f"{avg_comp}%"]

        hdr_cells = dash_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        val_cells = dash_table.rows[1].cells
        for i, v in enumerate(vals):
            val_cells[i].text = v
            val_cells[i].paragraphs[0].runs[0].font.size = Pt(14)
            val_cells[i].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)

        # Compliance Table
        doc.add_heading("Regulatory Compliance Posture", level=1)
        comp_table = doc.add_table(rows=1, cols=4)
        comp_hdr = comp_table.rows[0].cells
        comp_hdr[0].text = "Regulatory Framework"
        comp_hdr[1].text = "Violations"
        comp_hdr[2].text = "Compliance Score"
        comp_hdr[3].text = "Status"

        for fw, d in report.compliance_summary.items():
            row_cells = comp_table.add_row().cells
            row_cells[0].text = fw
            row_cells[1].text = str(d['violations_count'])
            row_cells[2].text = f"{d['compliance_score']}%"
            row_cells[3].text = d['status']

        doc.add_paragraph()

        # Detailed Findings Section
        doc.add_heading(f"Detailed Vulnerability Findings ({report.total_vulnerabilities})", level=1)

        for idx, vuln in enumerate(report.get_all_findings(), start=1):
            vuln_heading = doc.add_heading(f"#{idx}. {vuln.name} [{vuln.severity.value.upper()}]", level=2)
            
            p_details = doc.add_paragraph()
            p_details.add_run(f"Location: {vuln.url or vuln.host or 'N/A'}\n").bold = True
            p_details.add_run(f"CVSS Score: {vuln.cvss_score}   |   SLA Window: {vuln.priority}\n")
            p_details.add_run(f"OWASP 2025: {vuln.owasp_category}\n")
            p_details.add_run(f"Description: {vuln.description}\n")

            if vuln.proof_of_concept:
                doc.add_heading("Proof of Concept / Evidence", level=3)
                p_poc = doc.add_paragraph()
                r_poc = p_poc.add_run(vuln.proof_of_concept[:500])
                r_poc.font.name = 'Consolas'
                r_poc.font.size = Pt(8.5)

            doc.add_heading("Remediation Recommendation", level=3)
            doc.add_paragraph(vuln.remediation)

            doc.add_paragraph("─" * 40)

        doc.save(output_path)
        return output_path
